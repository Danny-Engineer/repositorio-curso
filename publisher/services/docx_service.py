from pathlib import Path
from docx import Document
from docx.shared import Pt
from config.settings import OUTPUT

def markdown_to_simple_docx(md_path: Path) -> Path:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    for line in md_path.read_text(encoding="utf-8").splitlines():
        clean = line.strip()
        if not clean or clean == "---":
            continue
        if clean.startswith("# "):
            doc.add_heading(clean[2:].strip(), level=1)
        elif clean.startswith("## "):
            doc.add_heading(clean[3:].strip(), level=2)
        elif clean.startswith("### "):
            doc.add_heading(clean[4:].strip(), level=3)
        elif clean.startswith("- "):
            doc.add_paragraph(clean[2:].strip(), style="List Bullet")
        elif clean.startswith("```"):
            continue
        else:
            doc.add_paragraph(clean)

    out = OUTPUT / f"{md_path.stem}.docx"
    doc.save(out)
    return out
